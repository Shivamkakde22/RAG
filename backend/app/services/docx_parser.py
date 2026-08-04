import docx

def extract_text(file_path: str):
    document = docx.Document(file_path)

    paragraphs = [p.text for p in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)

    return "\n".join(paragraphs)
